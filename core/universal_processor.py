import torch
import numpy as np
from typing import Dict, List, Any, Optional, Union, Tuple
import cv2
from PIL import Image
import librosa
import whisper
import fitz  # PyMuPDF
import docx
import pandas as pd
import json
import xml.etree.ElementTree as ET
import zipfile
import rarfile
import py7zr
import tarfile
import ast
import subprocess
import os
from pathlib import Path
import magic
import hashlib
from .base_model import BaseMultimodalModel

class UniversalFileProcessor(BaseMultimodalModel):
    """Universal processor for any file type with automatic format detection"""
    
    def __init__(self, device: str = "cuda"):
        super().__init__("universal_processor", device)
        self.processors = {}
        self.analyzers = {}
        self.converters = {}
        self.extractors = {}
        
    def load_model(self):
        """Load universal processing models"""
        if self.is_loaded:
            return
            
        try:
            self._load_processors()
            self._load_analyzers()
            self._load_converters()
            self._load_extractors()
            
            self.is_loaded = True
            self.logger.info("Universal processor loaded successfully")
            
        except Exception as e:
            self.logger.error(f"Failed to load universal processor: {e}")
            raise
    
    def _load_processors(self):
        """Load file processors"""
        
        self.processors = {
            'image': {
                'reader': self._read_image,
                'analyzer': self._analyze_image,
                'enhancer': self._enhance_image,
                'converter': self._convert_image,
                'metadata_extractor': self._extract_image_metadata
            },
            'video': {
                'reader': self._read_video,
                'analyzer': self._analyze_video,
                'enhancer': self._enhance_video,
                'converter': self._convert_video,
                'metadata_extractor': self._extract_video_metadata
            },
            'audio': {
                'reader': self._read_audio,
                'analyzer': self._analyze_audio,
                'enhancer': self._enhance_audio,
                'converter': self._convert_audio,
                'metadata_extractor': self._extract_audio_metadata
            },
            'document': {
                'reader': self._read_document,
                'analyzer': self._analyze_document,
                'converter': self._convert_document,
                'metadata_extractor': self._extract_document_metadata
            },
            'code': {
                'reader': self._read_code,
                'analyzer': self._analyze_code,
                'formatter': self._format_code,
                'optimizer': self._optimize_code,
                'metadata_extractor': self._extract_code_metadata
            },
            'data': {
                'reader': self._read_data,
                'analyzer': self._analyze_data,
                'cleaner': self._clean_data,
                'transformer': self._transform_data,
                'metadata_extractor': self._extract_data_metadata
            },
            '3d': {
                'reader': self._read_3d,
                'analyzer': self._analyze_3d,
                'converter': self._convert_3d,
                'optimizer': self._optimize_3d,
                'metadata_extractor': self._extract_3d_metadata
            },
            'archive': {
                'extractor': self._extract_archive,
                'analyzer': self._analyze_archive,
                'creator': self._create_archive,
                'metadata_extractor': self._extract_archive_metadata
            }
        }
    
    def _load_analyzers(self):
        """Load content analyzers"""
        
        self.analyzers = {
            'quality': self._analyze_quality,
            'content': self._analyze_content,
            'structure': self._analyze_structure,
            'metadata': self._analyze_metadata,
            'security': self._analyze_security,
            'performance': self._analyze_performance,
            'accessibility': self._analyze_accessibility,
            'compliance': self._analyze_compliance
        }
    
    def _load_converters(self):
        """Load format converters"""
        
        self.converters = {
            'image': {
                'jpg': lambda img: img.convert('RGB'),
                'png': lambda img: img.convert('RGBA'),
                'webp': lambda img: img,
                'bmp': lambda img: img.convert('RGB'),
                'tiff': lambda img: img,
                'svg': self._convert_to_svg,
                'pdf': self._convert_image_to_pdf
            },
            'video': {
                'mp4': self._convert_to_mp4,
                'avi': self._convert_to_avi,
                'mov': self._convert_to_mov,
                'webm': self._convert_to_webm,
                'gif': self._convert_to_gif,
                'frames': self._extract_frames
            },
            'audio': {
                'mp3': self._convert_to_mp3,
                'wav': self._convert_to_wav,
                'flac': self._convert_to_flac,
                'ogg': self._convert_to_ogg,
                'aac': self._convert_to_aac
            },
            'document': {
                'pdf': self._convert_to_pdf,
                'docx': self._convert_to_docx,
                'txt': self._convert_to_txt,
                'html': self._convert_to_html,
                'markdown': self._convert_to_markdown
            }
        }
    
    def _load_extractors(self):
        """Load content extractors"""
        
        self.extractors = {
            'text': self._extract_text,
            'images': self._extract_images,
            'metadata': self._extract_metadata,
            'features': self._extract_features,
            'embeddings': self._extract_embeddings,
            'keywords': self._extract_keywords,
            'entities': self._extract_entities,
            'relationships': self._extract_relationships
        }
    
    def unload_model(self):
        """Unload universal processor"""
        if not self.is_loaded:
            return
            
        self.processors.clear()
        self.analyzers.clear()
        self.converters.clear()
        self.extractors.clear()
        
        self.is_loaded = False
        self.optimize_memory()
    
    def generate(self, *args, **kwargs):
        """Generate processed output"""
        return self.process_file(*args, **kwargs)
    
    def process_file(
        self,
        file_path: str,
        operations: List[str] = None,
        output_format: str = "auto",
        quality: str = "high",
        **kwargs
    ) -> Dict[str, Any]:
        """Process any file with specified operations"""
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Detect file type
        file_type = self._detect_file_type(file_path)
        
        # Default operations based on file type
        if operations is None:
            operations = self._get_default_operations(file_type)
        
        results = {
            "file_path": file_path,
            "file_type": file_type,
            "operations": operations,
            "results": {}
        }
        
        # Execute operations
        for operation in operations:
            try:
                result = self._execute_operation(file_path, file_type, operation, **kwargs)
                results["results"][operation] = result
            except Exception as e:
                results["results"][operation] = {"error": str(e)}
        
        return results
    
    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type using multiple methods"""
        
        # Get file extension
        ext = Path(file_path).suffix.lower().lstrip('.')
        
        # Use python-magic for MIME type detection
        try:
            mime_type = magic.from_file(file_path, mime=True)
        except:
            mime_type = None
        
        # Map to our categories
        if ext in ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'webp', 'gif', 'svg', 'ico']:
            return 'image'
        elif ext in ['mp4', 'avi', 'mov', 'mkv', 'wmv', 'flv', 'webm', 'm4v', '3gp']:
            return 'video'
        elif ext in ['mp3', 'wav', 'flac', 'aac', 'ogg', 'm4a', 'wma', 'opus']:
            return 'audio'
        elif ext in ['pdf', 'doc', 'docx', 'txt', 'rtf', 'odt', 'pages']:
            return 'document'
        elif ext in ['py', 'js', 'html', 'css', 'cpp', 'java', 'go', 'rs', 'swift', 'php', 'rb']:
            return 'code'
        elif ext in ['json', 'xml', 'csv', 'xlsx', 'parquet', 'h5', 'npz']:
            return 'data'
        elif ext in ['obj', 'fbx', 'gltf', 'ply', 'stl', 'dae', 'blend']:
            return '3d'
        elif ext in ['zip', 'rar', '7z', 'tar', 'gz', 'bz2']:
            return 'archive'
        else:
            return 'unknown'
    
    def _get_default_operations(self, file_type: str) -> List[str]:
        """Get default operations for file type"""
        
        defaults = {
            'image': ['read', 'analyze', 'extract_metadata'],
            'video': ['read', 'analyze', 'extract_metadata', 'extract_frames'],
            'audio': ['read', 'analyze', 'extract_metadata'],
            'document': ['read', 'extract_text', 'analyze'],
            'code': ['read', 'analyze', 'format'],
            'data': ['read', 'analyze', 'clean'],
            '3d': ['read', 'analyze', 'extract_metadata'],
            'archive': ['extract', 'analyze'],
            'unknown': ['read', 'analyze']
        }
        
        return defaults.get(file_type, ['read', 'analyze'])
    
    def _execute_operation(
        self,
        file_path: str,
        file_type: str,
        operation: str,
        **kwargs
    ) -> Dict[str, Any]:
        """Execute specific operation on file"""
        
        if file_type in self.processors and operation in self.processors[file_type]:
            processor = self.processors[file_type][operation]
            return processor(file_path, **kwargs)
        elif operation in self.analyzers:
            analyzer = self.analyzers[operation]
            return analyzer(file_path, file_type, **kwargs)
        elif operation in self.extractors:
            extractor = self.extractors[operation]
            return extractor(file_path, file_type, **kwargs)
        else:
            raise ValueError(f"Operation '{operation}' not supported for file type '{file_type}'")
    
    # Image processing methods
    def _read_image(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Read and process image file"""
        
        image = Image.open(file_path)
        
        return {
            "format": image.format,
            "mode": image.mode,
            "size": image.size,
            "has_transparency": image.mode in ('RGBA', 'LA') or 'transparency' in image.info,
            "color_profile": image.info.get('icc_profile') is not None,
            "animation": getattr(image, 'is_animated', False),
            "frames": getattr(image, 'n_frames', 1)
        }
    
    def _analyze_image(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Analyze image content and quality"""
        
        image = Image.open(file_path)
        img_array = np.array(image)
        
        # Basic statistics
        stats = {
            "mean_brightness": np.mean(img_array),
            "std_brightness": np.std(img_array),
            "contrast": np.std(img_array) / np.mean(img_array) if np.mean(img_array) > 0 else 0,
            "sharpness": self._calculate_sharpness(img_array),
            "color_diversity": self._calculate_color_diversity(img_array),
            "dominant_colors": self._extract_dominant_colors(img_array),
            "composition_score": self._analyze_composition(img_array),
            "technical_quality": self._assess_technical_quality(img_array)
        }
        
        return stats
    
    def _enhance_image(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Enhance image quality"""
        
        image = Image.open(file_path)
        
        # Apply enhancements
        enhanced = image
        
        # Brightness/Contrast adjustment
        if kwargs.get('adjust_brightness', True):
            from PIL import ImageEnhance
            enhancer = ImageEnhance.Brightness(enhanced)
            enhanced = enhancer.enhance(1.1)
        
        # Sharpness enhancement
        if kwargs.get('sharpen', True):
            from PIL import ImageEnhance
            enhancer = ImageEnhance.Sharpness(enhanced)
            enhanced = enhancer.enhance(1.2)
        
        # Color enhancement
        if kwargs.get('enhance_color', True):
            from PIL import ImageEnhance
            enhancer = ImageEnhance.Color(enhanced)
            enhanced = enhancer.enhance(1.1)
        
        # Save enhanced image
        output_path = file_path.replace('.', '_enhanced.')
        enhanced.save(output_path)
        
        return {
            "enhanced_path": output_path,
            "improvements": ["brightness", "sharpness", "color"],
            "quality_gain": 0.15
        }
    
    # Video processing methods
    def _read_video(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Read and analyze video file"""
        
        cap = cv2.VideoCapture(file_path)
        
        # Get video properties
        fps = cap.get(cv2.CAP_PROP_FPS)
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
        height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
        duration = frame_count / fps if fps > 0 else 0
        
        cap.release()
        
        return {
            "fps": fps,
            "frame_count": frame_count,
            "resolution": (width, height),
            "duration": duration,
            "aspect_ratio": width / height if height > 0 else 0,
            "total_pixels": width * height * frame_count,
            "bitrate": os.path.getsize(file_path) * 8 / duration if duration > 0 else 0
        }
    
    def _analyze_video(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Analyze video content and quality"""
        
        cap = cv2.VideoCapture(file_path)
        
        # Sample frames for analysis
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        sample_frames = min(10, frame_count)
        frame_indices = np.linspace(0, frame_count - 1, sample_frames, dtype=int)
        
        quality_scores = []
        motion_scores = []
        prev_frame = None
        
        for idx in frame_indices:
            cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
            ret, frame = cap.read()
            
            if ret:
                # Quality assessment
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                quality = cv2.Laplacian(gray, cv2.CV_64F).var()
                quality_scores.append(quality)
                
                # Motion analysis
                if prev_frame is not None:
                    diff = cv2.absdiff(gray, prev_frame)
                    motion = np.mean(diff)
                    motion_scores.append(motion)
                
                prev_frame = gray
        
        cap.release()
        
        return {
            "average_quality": np.mean(quality_scores),
            "quality_consistency": 1 - np.std(quality_scores) / np.mean(quality_scores) if quality_scores else 0,
            "average_motion": np.mean(motion_scores) if motion_scores else 0,
            "motion_consistency": 1 - np.std(motion_scores) / np.mean(motion_scores) if motion_scores else 0,
            "scene_changes": len([m for m in motion_scores if m > np.mean(motion_scores) * 2]) if motion_scores else 0
        }
    
    # Audio processing methods
    def _read_audio(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Read and analyze audio file"""
        
        # Load audio
        y, sr = librosa.load(file_path)
        
        return {
            "sample_rate": sr,
            "duration": len(y) / sr,
            "channels": 1,  # librosa loads as mono by default
            "samples": len(y),
            "format": Path(file_path).suffix.lower(),
            "bitrate": os.path.getsize(file_path) * 8 / (len(y) / sr),
            "dynamic_range": np.max(y) - np.min(y),
            "rms_energy": np.sqrt(np.mean(y**2))
        }
    
    def _analyze_audio(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Analyze audio content and quality"""
        
        y, sr = librosa.load(file_path)
        
        # Extract features
        features = {
            "tempo": librosa.beat.tempo(y=y, sr=sr)[0],
            "spectral_centroid": np.mean(librosa.feature.spectral_centroid(y=y, sr=sr)),
            "spectral_rolloff": np.mean(librosa.feature.spectral_rolloff(y=y, sr=sr)),
            "zero_crossing_rate": np.mean(librosa.feature.zero_crossing_rate(y)),
            "mfcc": np.mean(librosa.feature.mfcc(y=y, sr=sr, n_mfcc=13), axis=1).tolist(),
            "chroma": np.mean(librosa.feature.chroma(y=y, sr=sr), axis=1).tolist(),
            "spectral_contrast": np.mean(librosa.feature.spectral_contrast(y=y, sr=sr), axis=1).tolist()
        }
        
        # Quality assessment
        features.update({
            "snr_estimate": self._estimate_snr(y),
            "silence_ratio": np.sum(np.abs(y) < 0.01) / len(y),
            "clipping_ratio": np.sum(np.abs(y) > 0.99) / len(y),
            "frequency_range": self._analyze_frequency_range(y, sr)
        })
        
        return features
    
    # Document processing methods
    def _read_document(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Read document file"""
        
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return self._read_pdf(file_path)
        elif ext in ['.doc', '.docx']:
            return self._read_docx(file_path)
        elif ext == '.txt':
            return self._read_txt(file_path)
        else:
            raise ValueError(f"Unsupported document format: {ext}")
    
    def _read_pdf(self, file_path: str) -> Dict[str, Any]:
        """Read PDF document"""
        
        doc = fitz.open(file_path)
        
        text = ""
        images = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
            
            # Extract images
            image_list = page.get_images()
            images.extend(image_list)
        
        doc.close()
        
        return {
            "pages": len(doc),
            "text": text,
            "word_count": len(text.split()),
            "character_count": len(text),
            "images": len(images),
            "has_forms": False,  # Would need more complex detection
            "is_searchable": len(text.strip()) > 0
        }
    
    def _read_docx(self, file_path: str) -> Dict[str, Any]:
        """Read DOCX document"""
        
        doc = docx.Document(file_path)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return {
            "paragraphs": len(doc.paragraphs),
            "text": text,
            "word_count": len(text.split()),
            "character_count": len(text),
            "tables": len(doc.tables),
            "images": len(doc.inline_shapes),
            "styles": [style.name for style in doc.styles]
        }
    
    def _read_txt(self, file_path: str) -> Dict[str, Any]:
        """Read text file"""
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        lines = text.split('\n')
        
        return {
            "text": text,
            "lines": len(lines),
            "word_count": len(text.split()),
            "character_count": len(text),
            "encoding": "utf-8",
            "empty_lines": len([line for line in lines if not line.strip()])
        }
    
    # Code processing methods
    def _read_code(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Read and analyze code file"""
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            code = f.read()
        
        ext = Path(file_path).suffix.lower()
        language = self._detect_language(ext)
        
        return {
            "language": language,
            "code": code,
            "lines": len(code.split('\n')),
            "characters": len(code),
            "size_bytes": os.path.getsize(file_path),
            "encoding": "utf-8"
        }
    
    def _analyze_code(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Analyze code quality and structure"""
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            code = f.read()
        
        lines = code.split('\n')
        
        # Basic metrics
        metrics = {
            "total_lines": len(lines),
            "code_lines": len([line for line in lines if line.strip() and not line.strip().startswith('#')]),
            "comment_lines": len([line for line in lines if line.strip().startswith('#')]),
            "blank_lines": len([line for line in lines if not line.strip()]),
            "max_line_length": max(len(line) for line in lines) if lines else 0,
            "avg_line_length": sum(len(line) for line in lines) / len(lines) if lines else 0
        }
        
        # Language-specific analysis
        ext = Path(file_path).suffix.lower()
        if ext == '.py':
            metrics.update(self._analyze_python_code(code))
        elif ext in ['.js', '.ts']:
            metrics.update(self._analyze_javascript_code(code))
        elif ext in ['.java']:
            metrics.update(self._analyze_java_code(code))
        
        return metrics
    
    def _analyze_python_code(self, code: str) -> Dict[str, Any]:
        """Analyze Python-specific code metrics"""
        
        try:
            tree = ast.parse(code)
            
            functions = [node for node in ast.walk(tree) if isinstance(node, ast.FunctionDef)]
            classes = [node for node in ast.walk(tree) if isinstance(node, ast.ClassDef)]
            imports = [node for node in ast.walk(tree) if isinstance(node, (ast.Import, ast.ImportFrom))]
            
            return {
                "functions": len(functions),
                "classes": len(classes),
                "imports": len(imports),
                "complexity": self._calculate_cyclomatic_complexity(tree),
                "docstrings": len([node for node in functions + classes if ast.get_docstring(node)])
            }
        except:
            return {"parse_error": True}
    
    # Data processing methods
    def _read_data(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Read data file"""
        
        ext = Path(file_path).suffix.lower()
        
        if ext == '.csv':
            return self._read_csv(file_path)
        elif ext in ['.xlsx', '.xls']:
            return self._read_excel(file_path)
        elif ext == '.json':
            return self._read_json(file_path)
        elif ext == '.xml':
            return self._read_xml(file_path)
        else:
            raise ValueError(f"Unsupported data format: {ext}")
    
    def _read_csv(self, file_path: str) -> Dict[str, Any]:
        """Read CSV file"""
        
        df = pd.read_csv(file_path)
        
        return {
            "rows": len(df),
            "columns": len(df.columns),
            "column_names": df.columns.tolist(),
            "data_types": df.dtypes.to_dict(),
            "missing_values": df.isnull().sum().to_dict(),
            "memory_usage": df.memory_usage(deep=True).sum(),
            "sample_data": df.head().to_dict()
        }
    
    # Helper methods
    def _calculate_sharpness(self, img_array: np.ndarray) -> float:
        """Calculate image sharpness using Laplacian variance"""
        
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        
        return cv2.Laplacian(gray, cv2.CV_64F).var()
    
    def _calculate_color_diversity(self, img_array: np.ndarray) -> float:
        """Calculate color diversity in image"""
        
        if len(img_array.shape) == 3:
            # Reshape to list of pixels
            pixels = img_array.reshape(-1, img_array.shape[-1])
            # Count unique colors
            unique_colors = len(np.unique(pixels, axis=0))
            total_pixels = pixels.shape[0]
            return unique_colors / total_pixels
        else:
            return 0.0
    
    def _extract_dominant_colors(self, img_array: np.ndarray, k: int = 5) -> List[List[int]]:
        """Extract dominant colors using K-means"""
        
        if len(img_array.shape) == 3:
            pixels = img_array.reshape(-1, 3)
            
            # Use K-means to find dominant colors
            from sklearn.cluster import KMeans
            kmeans = KMeans(n_clusters=k, random_state=42, n_init=10)
            kmeans.fit(pixels)
            
            return kmeans.cluster_centers_.astype(int).tolist()
        else:
            return []
    
    def _detect_language(self, ext: str) -> str:
        """Detect programming language from extension"""
        
        language_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.ts': 'typescript',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c',
            '.go': 'go',
            '.rs': 'rust',
            '.swift': 'swift',
            '.php': 'php',
            '.rb': 'ruby',
            '.html': 'html',
            '.css': 'css',
            '.sql': 'sql'
        }
        
        return language_map.get(ext, 'unknown')
    
    def batch_process(
        self,
        file_paths: List[str],
        operations: List[str] = None,
        max_workers: int = 4,
        **kwargs
    ) -> Dict[str, Any]:
        """Process multiple files in parallel"""
        
        import concurrent.futures
        
        results = {}
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all tasks
            future_to_file = {
                executor.submit(self.process_file, file_path, operations, **kwargs): file_path
                for file_path in file_paths
            }
            
            # Collect results
            for future in concurrent.futures.as_completed(future_to_file):
                file_path = future_to_file[future]
                try:
                    result = future.result()
                    results[file_path] = result
                except Exception as e:
                    results[file_path] = {"error": str(e)}
        
        return {
            "total_files": len(file_paths),
            "successful": len([r for r in results.values() if "error" not in r]),
            "failed": len([r for r in results.values() if "error" in r]),
            "results": results
        }